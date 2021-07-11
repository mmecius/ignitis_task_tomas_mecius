import os
from os.path import basename
from zipfile import ZipFile
from read_excel_file import ExcelFileManager
import requests
from selenium import webdriver  # type: ignore
from time import sleep
from docx import Document  # type: ignore
from docx.shared import Inches  # type: ignore
import pafy  # type: ignore
from send_email import send_email_with_files
import glob

read_excel_file = ExcelFileManager()


def search_youtube():
    """Opens Youtube and starts searching."""
    search_url = "https://www.youtube.com/results?search_query="
    driver.get(f"{search_url}{group_name}")


def close_pop():
    """Closes pop up"""
    pop_up_i_agree = driver.find_element_by_xpath(
        '//*[@id="yDmH0d"]/c-wiz/div/div/div/div[2]/div[1]/div[4]/form/div[1]/div/button'
    )
    pop_up_i_agree.click()


def select_filter_button():
    """In Youtube search selects filter button"""
    filter_button = driver.find_element_by_xpath(
        '//*[@id="container"]/ytd-toggle-button-renderer/a'
    )
    filter_button.click()


def select_view_count():
    """In filter selects View Count as sorting parameter."""
    filter_by_rating = driver.find_element_by_xpath(
        "/html/body/ytd-app/div/ytd-page-manager/ytd-search/div[1]/ytd-two-column-search-results-renderer/div/ytd-section-list-renderer/div[1]/div[2]/ytd-search-sub-menu-renderer/div[1]/iron-collapse/div/ytd-search-filter-group-renderer[5]/ytd-search-filter-renderer[3]/a/div/yt-formatted-string"
    )
    filter_by_rating.click()


def select_filter_time():
    """In filter selects time filter."""
    filter_4_20_min = driver.find_element_by_xpath(
        "/html/body/ytd-app/div/ytd-page-manager/ytd-search/div[1]/ytd-two-column-search-results-renderer/div/ytd-section-list-renderer/div[1]/div[2]/ytd-search-sub-menu-renderer/div[1]/iron-collapse/div/ytd-search-filter-group-renderer[3]/ytd-search-filter-renderer[2]/a/div/yt-formatted-string"
    )
    filter_4_20_min.click()


html = """\
    <html>
      <body>
        <p>Hello,<br>
           Please find attached Music Groups scrapped from Youtube:<br>
        </p>
        <p>Best wishes,<br>
        Robot</p>
      </body>
    </html>
    """


def save_img_to_folder(url, group_name, count):
    """From youtube URL extracts .jpg image and saves it to files/img/ directory as *.jpg"""
    response = requests.get(url)
    with open(f"files/img/{group_name}{count}.jpg", "wb") as file:
        file.write(response.content)
    saved_img = f"{group_name}{count}.jpg"
    return saved_img


def write_to_doc_file(group_name, title, image):
    """Writes to *.docx file. File name as group_name. Title and image to *.docx file."""
    document = Document()
    document.add_heading(title, 0)
    document.add_picture(f"files/img/{image}", width=Inches(4.0), height=Inches(4.0))
    document.save(f"files/doc_files/{group_name}.docx")


def zip_files(group_name):
    """Zips *.docx files in files/doc_files/ directory"""
    dirName = "files/doc_files/"
    with ZipFile(f"{group_name}.zip", "w") as zipObj:
        for folderName, subfolders, filenames in os.walk(dirName):
            for filename in filenames:
                filePath = os.path.join(folderName, filename)
                zipObj.write(filePath, basename(filePath))


def remove_doc_files(doc_files):
    """Removes *.docx files"""
    for file in doc_files:
        os.remove(file)


def remove_img_files(img_files):
    """Removes *.img files"""
    for file in img_files:
        os.remove(file)


def remove_zip_files(zip_files):
    """Removes *.zip files"""
    for file in zip_files:
        os.remove(file)


# Functions

list_groups = read_excel_file.read_excel_file_to_list()
for group_name in list_groups:
    print(group_name)
    url = "https://www.youtube.com"
    chrome_driver_path = "/Users/tomasmecius/Development/chromedriver"
    driver = webdriver.Chrome(executable_path=chrome_driver_path)
    search_youtube()
    close_pop()
    sleep(2)
    select_filter_button()
    sleep(2)
    select_view_count()
    sleep(2)
    select_filter_button()
    sleep(2)
    select_filter_time()
    sleep(2)
    user_data = driver.find_elements_by_xpath('//*[@id="video-title"]')
    links = []

    for i in user_data:
        links.append(i.get_attribute("href"))

    filtered_links = []
    for i in links:
        if i != None:
            try:
                video = pafy.new(i)
            except KeyError:
                continue
            if video.length <= 360:
                filtered_links.append(i)

    driver.close()
    driver.quit()

    for url in filtered_links[:3]:
        count = +1
        video = pafy.new(url)
        title = video.title
        img_url = video.thumb
        image = save_img_to_folder(img_url, group_name, count)
        write_to_doc_file(group_name, title, image)

    zip_files(group_name)


doc_files = glob.glob("files/doc_files/*.docx")
img_files = glob.glob("files/img/*.jpg")
zip_files = glob.glob("*.zip")
send_email_with_files(html, zip_files)
remove_img_files(img_files)
remove_doc_files(doc_files)
remove_zip_files(zip_files)
