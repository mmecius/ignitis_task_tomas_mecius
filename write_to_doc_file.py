from docx import Document  # type: ignore
from docx.shared import Inches  # type: ignore


class DocFileManager:
    def __init__(self, title, image, group_name):
        self.title = title
        self.image = image
        self.group_name = group_name

    def write_to_doc_file(self, group_name, title, image):
        """Writes url data to doc."""
        document = Document()
        document.add_heading(title, 0)
        document.add_picture(image, width=Inches(4.0), height=Inches(4.0))
        document.save(f"files/doc_files/{group_name}.docx")
